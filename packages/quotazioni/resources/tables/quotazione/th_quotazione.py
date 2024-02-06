#!/usr/bin/python3
# -*- coding: utf-8 -*-

from gnr.web.gnrbaseclasses import BaseComponent
from gnr.core.gnrdecorator import public_method
from gnr.web.gnrbaseclasses import TableTemplateToHtml
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches
import base64
from bs4 import BeautifulSoup

class View(BaseComponent):

    def th_struct(self,struct):
        r = struct.view().rows()
        r.fieldcell('data')
        r.fieldcell('quot_n', width='8em')
        r.fieldcell('oggetto', width='50%')
        r.fieldcell('cliente_id', width='100%')

    def th_order(self):
        return 'data'

    def th_query(self):
        return dict(column='full_quot', op='contains', val='')

    def th_options(self):
        return dict(partitioned=True)

class Form(BaseComponent):
    py_requires='gnrcomponents/pagededitor/pagededitor:PagedEditor,gnrcomponents/attachmanager/attachmanager:AttachManager'

    def th_form(self, form):
        tc = form.center.tabContainer()
        bc = tc.borderContainer(title='!![en]Quotation')
        bc_att = tc.borderContainer(title='!![en]<strong>Quotation received</strong>')
       # bc_att.borderContainer(region='center',height='auto', background = '#f2f0e8', splitter=True).contentPane(title='!![en]Quotation Received',pageName='quot_rec',height='100%').remote(self.quot_recLazyMode,_waitingMessage='!![en]Please wait')
        self.quotazioneRicevuta(bc_att.contentPane(title='!![en]Quotation received',pageName='quot_rec',height='100%'))
        self.datiQuotazione(bc.borderContainer(region='top',datapath='.record',height='130px', splitter=True))
        #self.datiQuotazione(bc.roundedGroupFrame(title='!![en]Quotation',region='top',datapath='.record',height='130px', background='lightgrey', splitter=True))
        tc = bc.tabContainer(region = 'center',margin='2px',selectedPage='^.tabname')
       #self.dettagli(tc.borderContainer(title='!![en]Proforma').contentPane(region='center',datapath='.record'))
        self.description(tc.contentPane(title='!![en]First description'))
        self.extra(tc.contentPane(title='!![en]Extra details'))
        self.servIncl(tc.contentPane(title='!![en]Services included'))
        self.surcharges(tc.contentPane(title='!![en]Surcharges'))
        self.specCond(tc.contentPane(title='!![en]Specific conditions'))
        self.servExcl(tc.contentPane(title='!![en]Services excluded'))
        self.workTimes(tc.contentPane(title='!![en]Working times'))
        self.note(tc.contentPane(title='!![en]Notes'))
        self.paymentCond(tc.contentPane(title='!![en]Payment conditions'))
        
        bc = bc.borderContainer(region='right', width='48%', splitter=True)#.borderContainer(region='top', height='50%')
        tc_right = bc.tabContainer(region='center',margin='2px')
        self.allegatiQuotazione(tc_right.contentPane(title='!![en]Attachments', height='100%'))
        self.editQuotation(tc_right.framePane(title='!![EN]Edit Quotation', datapath='#FORM.editPagine'))
    
   #@public_method
   #def quot_recLazyMode(self,pane):    
   #    pane.stackTableHandler(relation='@quot_ricevuta',view_store__onBuilt=True,liveUpdate=True)

    def quotazioneRicevuta(self,pane):
        pane.stackTableHandler(relation='@quot_ricevuta',viewResource='View')

    def dettagli(self,pane):
        #leftdett = pane.roundedGroup(title='!![en]Proforma details', height='auto')
        #fb = leftdett.formbuilder(cols=1, border_spacing='4px',margin='4px')
        fb = pane.formbuilder(cols=1, border_spacing='4px',margin='4px',region='left')
        fb.div('',width='60em')
        
       
        
        #frame.simpleTextArea(title='Body Quotation',value='^.body', editor=True)
    def description(self,pane):
        pane.inlineTableHandler(relation='@description_quot',viewResource='ViewFromDescription')
    
    def extra(self,pane):
        pane.inlineTableHandler(relation='@quotazione_extra',viewResource='ViewFromExtracost')
    
    def servIncl(self,pane):
        pane.inlineTableHandler(relation='@servincl_quot',viewResource='ViewFromServIncl')

    def servExcl(self,pane):
        pane.inlineTableHandler(relation='@serviexcl_quot',viewResource='ViewFromServExcl')

    def specCond(self,pane):
        pane.inlineTableHandler(relation='@specific_quot',viewResource='ViewFromSpecificCond')    

    def surcharges(self,pane):
        pane.inlineTableHandler(relation='@surcharge_quot',viewResource='ViewFromSurcharges') 

    def workTimes(self,pane):
        pane.inlineTableHandler(relation='@times_quot',viewResource='ViewFromTimesWork')

    def note(self,pane):
        pane.inlineTableHandler(relation='@note_quot',viewResource='ViewFromNote')

    def paymentCond(self,pane):
        pane.inlineTableHandler(relation='@paym_quot',viewResource='ViewFromPaymCond')            

    def allegatiQuotazione(self,pane):
        pane.attachmentGrid()

    def datiQuotazione(self,bc):
        bc.contentPane(region='center').linkerBox('cliente_id',margin='2px',openIfEmpty=True, validate_notnull=True,
                                                    columns='$rag_sociale,$address,$cap,$city,$email',
                                                    auxColumns='$tel',
                                                #    clientTemplate=True,
                                                    newRecordOnly=True,formResource='Form',dialog_windowRatio = 1)
                                                    #dialog_height='500px',dialog_width='800px')
        left = bc.roundedGroup(title='!![en]Quotation datas',region='left',width='45%')
        fb = left.formbuilder(cols=3, border_spacing='4px', width='90%')
        #fb = pane.div(margin_left='50px',margin_right='80px').formbuilder(cols=4, border_spacing='4px',fld_width='10em')
        
        fb.field('data' )
        fb.field('quot_n', readOnly=True, width='100%')
        fb.br()
        fb.field('oggetto', colspan=3, width='98%',height='100%', tag='simpleTextArea' )
        
        #fb.field('cliente_id' )
    def th_bottom_custom(self, bottom):
        bar = bottom.slotBar('10,stampa_quotazione,5,email_quotazione,5,stampa_quot_template,5,quotazione_doc,*,10')
        btn_quot_print=bar.stampa_quotazione.button('!![en]Print quotation',
                                                    action="""genro.publish("table_script_run",{table:"quotazioni.quotazione",
                                                                               res_type:'print',
                                                                               resource:'stampa_quotazione',
                                                                               pkey: pkey})
                                                                               """,
                                                                               pkey='=#FORM.pkey')
        bar.email_quotazione.button('!![en]Email quotation', iconClass='print',
                                        action="""genro.publish("table_script_run",{table:"quotazioni.quotazione",
                                                                                   res_type:'print',
                                                                                   resource:'email_quotazione',
                                                                                   pkey: pkey})""",
                                                                                   pkey='=#FORM.pkey')

        btn_quot_print=bar.stampa_quot_template.button('!![en]Print quotation template')
        btn_quot = bar.quotazione_doc.button('!![en]Quotanion doc.')
        btn_quot_print.dataRpc('nome_temp', self.print_quotation,record='=#FORM.record',nome_template = 'quotazioni.quotazione:quotation',format_page='A4')
        btn_quot.dataRpc('nome_temp', self.apridoc,record='=#FORM.record',nome_form='quotazione',agency_name='=#FORM.record.@agency_id.agency_name',
                         workport='=#FORM.record.@agency_id.@port.descrizione',_virtual_column='$cliente,$cliente_br')

    def editQuotation(self, frame):
        bar = frame.top.slotBar('10, lett_select,*',height='20px',border_bottom='1px solid silver')
        fb = bar.lett_select.formbuilder(cols=2,datapath='#FORM.record.htmlbag_quot')
        fb.dbselect('^.letterhead_id',table='adm.htmltemplate',lbl='carta intestata',hasDownArrow=True)
        fb.button('Get Html Quotation Doc').dataRpc('#FORM.record.htmlbag_quot.source',self.db.table('quotazioni.quotazione').getHTMLDoc,
                                            quot_id='=#FORM.pkey',
                                            record_template='quotation',
                                            letterhead='.letterhead_id')
        
        frame.pagedEditor(value='^#FORM.record.htmlbag_quot.source',pagedText='^#FORM.record.htmlbag_quot.output',
                          border='1px solid silver',
                          letterhead_id='^#FORM.record.htmlbag_quot.letterhead_id',
                          datasource='#FORM.record',printAction=True)
    
    @public_method
    def print_quotation(self, record, resultAttr=None, nome_template=None, format_page=None, **kwargs):
        #msg_special=None
        record_id=record['id']
        #print(x)
       #if selId is None:
       #    msg_special = 'yes'
       #    return msg_special

        tbl_sof = self.db.table('quotazioni.quotazione')
        builder = TableTemplateToHtml(table=tbl_sof)

        nome_temp = nome_template.replace('quotazioni.quotazione:','')
        nome_file = '{cl_id}.pdf'.format(
                    cl_id=nome_temp)

        template = self.loadTemplate(nome_template)  # nome del template
        pdfpath = self.site.storageNode('home:stampe_template', nome_file)

        #preleviamo l'id della carta intestata dalla table agency da passare nel builder di stampa
        ag_id = self.db.currentEnv.get('current_agency_id')
        tbl_agency =  self.db.table('agz.agency')
        htmltemplate_id = tbl_agency.readColumns(columns='htmltemplate_id',
                  where='$id=:ag_id',
                    ag_id=ag_id)
        #print(x)
        #builder(record=selId, template=template)
        builder(record=record_id, template=template, letterhead_id=htmltemplate_id)
        #selezioniamo il service di stampa
        #builder.pdf_service='wk'

        if format_page=='A3':
            builder.page_format='A3'
            builder.page_width=427
            builder.page_height=290

        result = builder.writePdf(pdfpath=pdfpath)
        #print(x)
        self.setInClientData(path='gnr.clientprint',
                              value=result.url(timestamp=datetime.now()), fired=True)
    
    @public_method
    def apridoc(self,record,nome_form=None, agency_name=None, workport=None, **kwargs):
        #prendiamo il percorso del file quotation.doc negli allegati dell'Agenzia
        id_agz_atc=record['agency_id']
        tbl_agzatt =  self.db.table('agz.agency_atc')
        fileurl = tbl_agzatt.query(columns='$fileurl,$description',
                  where='$maintable_id=:mt_id AND $description=:descr',mt_id=id_agz_atc, descr='quotation').fetch()
        file_url = fileurl[0]['fileurl']
        file_path = file_url.replace('/home','site')
        file_sn = self.site.storageNode(file_path)
        #print(x)
        date = record['data'].strftime("%d/%m/%Y")
        quot_n= record['quot_n']
        #verifichiamo se i record sono None e passiamo un valore '' per non avere l'errore di stringa durante la sostizione sul docx
        if record['oggetto'] is None:
            oggetto=''
        else:            
            oggetto=record['oggetto']
        #con BeautifulSoup andiamo a rimuove i tag html ai record che andremo a inserire nella doc    
        if record['descr_costo'] is None:
            descr_costo = ''
        else:    
            descr_costo = BeautifulSoup(record['descr_costo'], "html.parser").text
        if record['serv_included'] is None:
            serv_included = ''
        else:            
            serv_included= BeautifulSoup(record['serv_included'], "html.parser").text
        if record['serv_excluded']is None:
            serv_excluded = ''
        else:
            serv_excluded= BeautifulSoup(record['serv_excluded'], "html.parser").text
        if record['special_cond'] is None:
            special_cond = ''
        else:    
            special_cond= BeautifulSoup(record['special_cond'], "html.parser").text
        if record['times_work'] is None:
            times_work = ''
        else:    
            times_work= BeautifulSoup(record['times_work'], "html.parser").text
        if record['note'] is None:
            note = ''
        else:    
            note= BeautifulSoup(record['note'], "html.parser").text
        #cliente=record['cliente_br']
        cliente= BeautifulSoup(record['cliente_br'], "html.parser").text
        agency_stamp=record['@agency_id.agency_stamp']
        img_bytes = base64.b64decode(agency_stamp.replace('data:image/png;base64,',''))
        image_stamp = io.BytesIO(img_bytes)
        #img_string = agency_stamp.replace('data:image/png;base64,','')
        #image_stamp = io.StringIO(img_string)
        #print(x)
        if nome_form=='quotazione':
            #tbl_place = self.db.table('unlocode.place')
            #san_place=tbl_place.readColumns(columns="$descrizione || ' - ' || @nazione_code.nome", where='$id=:place_id', place_id=san_place_id)
        
            nome_file = 'Quotation.docx'
            nome_file_out = 'Quotation_filled.docx'
            file_sn_out = self.site.storageNode('home:form_standard', 'Quotation_filled.docx')
           
            variables = {
            "${cliente}": cliente,    
            "${quot_n}": quot_n,
            "${oggetto}": oggetto,
            "${descr_costo}": descr_costo,
            "${serv_included}": serv_included,
            "${serv_excluded}": serv_excluded,
            "${special_cond}": special_cond,
            "${times_work}": times_work,
            "${note}": note,
            "${luogo}": workport,
            "${data}": date,
            "${nome_ditta}": agency_name,
            }
        
        #file_sn = self.site.storageNode('home:form_standard', nome_file)
        template_file_path = file_sn.internal_path
        #template_file_path = '/home/tommaso/Documenti/Linux/Python/ModificaDocx/test.docx'
        
        output_file_path = file_sn_out.internal_path
        #output_file_path = '/home/tommaso/Documenti/Linux/Python/ModificaDocx/result.docx'
        
        

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
        
        #troviamo la tabella che ci interessa nella doc e inseriamo il nostro timbro
        table1 = template_document.tables[1]
        col1 = table1.columns[1]
        cell = col1.cells[0]
        paragr = cell.paragraphs[0]
        run = paragr.add_run()
        run.add_picture(image_stamp, width=Inches(1.25))
        #qui troviamo la table che si trova nell'intestazione della doc se vogliamo inserire qualcosa
        #section = template_document.sections[0]
        #header = section.header
        #h_paragr = header.tables[0].columns[0].cells[0].paragraphs[0]
        #run = h_paragr.add_run()
        #run.add_picture(image_stamp, width=Inches(1.25))

        template_document.save(output_file_path)
        #apriamo direttamente il file salvato con il programma standard di sistema
        filename=output_file_path
        
       # subprocess.call(('xdg-open', filename))
        
      # path=self.site.site_path + str('/form_standard')
      # doc_path=filename
      # subprocess.call(['libreoffice',
      #          # '--headless',
      #          '--convert-to',
      #          'pdf',
      #          '--outdir',
      #          path,
      #          doc_path])
      # nome_file = nome_form + str('.pdf')  
              
        path_pdf = self.site.storageNode('home:form_standard', nome_file_out)
        
        #path_pdf=path + str('/') + nome_form + str('.pdf')
        result=self.site.storageNode(path_pdf)
        #print(x)
        self.setInClientData(path='gnr.clientprint',
                              value=result.url(), fired=True)
        #print(x)
        return nome_form

    def replace_text_in_paragraph(self,paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

    def th_options(self):
        return dict(dialog_height='400px', dialog_width='600px' )
